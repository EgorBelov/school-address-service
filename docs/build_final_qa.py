"""
Скрипт генерации docs/финальный_отчет_тестировщик.docx — объединённый
отчёт: общая часть проекта + индивидуальная часть тестировщика, в одном
документе с единой титульной страницей.

Принцип работы:
1) перегенерирует docs/общая_часть.docx и docs/индивидуальная_тестировщик.docx
   запуском build_general.py и build_qa.py — чтобы исходные части были
   актуальны;
2) создаёт новый документ с единым титульным листом;
3) копирует body-элементы общей части, начиная со следующего элемента
   после первого page break (то есть пропуская титул источника);
4) добавляет разделитель «ИНДИВИДУАЛЬНАЯ ЧАСТЬ — ТЕСТИРОВЩИК»;
5) копирует body-элементы индивидуальной части по тому же принципу.

Исходные файлы общая_часть.docx и индивидуальная_тестировщик.docx остаются
на месте — этот скрипт их НЕ удаляет.
"""
import subprocess
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from report_lib import (
    FONT,
    FONT_SIZE,
    LINE_SPACING,
    _set_font,
    add_title_page,
    new_document,
)


DOCS_DIR = Path(__file__).parent
GENERAL_PATH = DOCS_DIR / "общая_часть.docx"
INDIVIDUAL_PATH = DOCS_DIR / "индивидуальная_тестировщик.docx"
FINAL_PATH = DOCS_DIR / "финальный_отчет_тестировщик.docx"


def _find_first_page_break_body_index(doc):
    """Вернуть индекс body-элемента (параграфа), содержащего первый
    page break. Возвращает -1, если page break не найден."""
    body = doc.element.body
    for i, child in enumerate(body):
        for br in child.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return i
    return -1


def _copy_body_after_title(src_doc, dst_doc):
    """Скопировать в dst_doc все body-элементы src_doc, расположенные
    после первого page break (то есть пропуская титульную страницу).
    Элемент sectPr (свойства секции) не копируется — у dst_doc свой."""
    skip_until = _find_first_page_break_body_index(src_doc)
    if skip_until < 0:
        raise RuntimeError(
            f"В документе {src_doc} не найден page break, "
            "разделяющий титул и содержимое."
        )
    src_body_children = list(src_doc.element.body)
    for el in src_body_children[skip_until + 1:]:
        tag = el.tag.split("}")[-1]
        if tag == "sectPr":
            continue
        dst_doc.element.body.append(deepcopy(el))


def _add_part_divider(doc, title_text):
    """Большой центрированный заголовок-разделитель между частями отчёта."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = 0
    pf.space_after = 0

    # Несколько пустых параграфов сверху — чтобы заголовок оказался
    # примерно в верхней трети страницы.
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run(title_text)
    _set_font(run, size=FONT_SIZE + 4, bold=True)

    doc.add_page_break()


# ─── 1) Перегенерируем исходные части ─────────────────────────

print("[1/3] Перегенерация исходных частей…")
for script in ("build_general.py", "build_qa.py"):
    subprocess.run([sys.executable, str(DOCS_DIR / script)], check=True)


# ─── 2) Создаём финальный документ с единым титулом ──────────

print("[2/3] Сборка финального документа…")
final_doc = new_document()
add_title_page(
    final_doc,
    "Отчёт о групповом проекте.\nЧасть тестировщика "
    "(общая часть и индивидуальная часть)",
)


# ─── 3) Копируем общую часть, разделитель, индивидуальную часть ───

general_doc = Document(str(GENERAL_PATH))
_copy_body_after_title(general_doc, final_doc)

_add_part_divider(final_doc, "ИНДИВИДУАЛЬНАЯ ЧАСТЬ\nТЕСТИРОВЩИК")

individual_doc = Document(str(INDIVIDUAL_PATH))
_copy_body_after_title(individual_doc, final_doc)


# ─── Сохранение ───────────────────────────────────────────────

print("[3/3] Сохранение результата…")
final_doc.save(str(FINAL_PATH))
print(f"Готово: {FINAL_PATH}")
