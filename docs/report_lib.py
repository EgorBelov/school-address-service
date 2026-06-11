"""
Общая инфраструктура для генерации отчётов (общая часть + 6 индивидуальных)
в визуальном стиле ВШЭ.

Все стили (HSE Default Text, HSE Title 1/2/3, HSE list, Caption, поля
страницы, нумерация маркированных списков) берутся из шаблона
_template_hse.docx, который представляет собой эталон «отчет_Курсовая.docx»
с очищенным body. new_document() открывает этот шаблон, остальные функции
лишь применяют нужный style к параграфам.

Кратко об используемых стилях:
  • HSE Default Text — основной текст, TNR 13pt, 1.5 межстрочный, justify,
    отступ первой строки 1.25 см (709 twips).
  • HSE Title 1 — заголовок главы. TNR 16pt bold, по центру, pageBreakBefore.
    Поэтому НЕ нужно вручную вызывать doc.add_page_break() перед главой.
  • HSE Title 1 NoTOC — то же, но не попадает в оглавление (введение,
    заключение, аннотация).
  • HSE Title 2 — подраздел. TNR 14pt, по центру (наследует от Title 1).
  • HSE Title 3 — подподраздел. TNR 13pt.
  • HSE list — маркированный список с длинным тире в качестве буллета.
  • Caption — подпись рисунка/таблицы, italic 9pt, цветом text2.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


TEMPLATE_PATH = Path(__file__).parent / "_template_hse.docx"

# Константы для редких случаев, когда шрифт надо задать вручную
# (например, divider в build_final_qa.py или текст на титульном листе).
FONT = "Times New Roman"
FONT_SIZE = 13
LINE_SPACING = 1.5


def _set_font(run, name="Times New Roman", size=None, bold=False, italic=False):
    """Задаёт шрифт run явно — нужно только для титульной страницы, где
    мы хотим конкретные размеры независимо от стилей. В обычном тексте
    шрифт берётся из стиля HSE Default Text."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading(doc, text, level=1, notoc=False):
    """Заголовок главы/подраздела/подподраздела.
    notoc=True для введения/заключения/аннотации — стиль HSE Title 1 NoTOC."""
    if level == 1:
        style = "HSE Title 1 NoTOC" if notoc else "HSE Title 1"
    elif level == 2:
        style = "HSE Title 2"
    else:
        style = "HSE Title 3"
    p = doc.add_paragraph(text, style=style)
    return p


def add_para(doc, text, bold=False, italic=False,
             first_line_indent=None, align=None):
    """Обычный абзац стилем HSE Default Text.
    first_line_indent / align — опциональные перекрытия (например, 0 для
    подписи или CENTER для центровки). Если bold/italic = True, форматирование
    применяется к run'у поверх стиля."""
    p = doc.add_paragraph(style="HSE Default Text")
    pf = p.paragraph_format
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if align is not None:
        pf.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_bullet(doc, text):
    """Маркированный список (стиль HSE list — буллет «—» из шаблона)."""
    p = doc.add_paragraph(text, style="HSE list")
    return p


def _set_row_header(row):
    """Помечает строку как заголовок таблицы (<w:tblHeader/>).
    При переносе таблицы на следующую страницу Word автоматически
    повторяет такие строки в качестве шапки."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def _set_row_cant_split(row):
    """Запрещает разрыв строки между страницами (<w:cantSplit/>).
    Если ячейка не помещается, вся строка переедет на новую страницу."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def add_table_simple(doc, rows, header=True, col_widths=None):
    """Простая таблица. rows[0] — шапка. Текст в ячейках — TNR 12pt,
    шапка bold; стили внутри ячеек переопределяются явно, чтобы не
    наследовать отступ первой строки от HSE Default Text.

    Шапка таблицы помечена <w:tblHeader/> — при переносе таблицы на
    новую страницу Word автоматически повторит эту строку. Каждая
    строка помечена <w:cantSplit/>, чтобы текст в ячейке не разрывался
    между страницами."""
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, row in enumerate(rows):
        tbl_row = table.rows[i]
        _set_row_cant_split(tbl_row)
        if header and i == 0:
            _set_row_header(tbl_row)
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(cell_text)
            _set_font(run, size=12, bold=(header and i == 0))

    if col_widths:
        for j, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(width)
    return table


def add_note(doc, text):
    """Курсивный блок-инструкция «📝 что писать в этом разделе» для
    скелетных индивидуальных шаблонов (build_lead/ba/sa/dev/ai).
    HSE-стиль для них не используется — это не часть финального текста."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)
    pf.space_before = Pt(2)
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("📝 ")
    _set_font(run, size=13, bold=True, italic=True)
    run2 = p.add_run(text)
    _set_font(run2, size=13, italic=True)
    return p


def _apply_caption_overrides(p, alignment, keep_next):
    """Воссоздаёт визуал подписи из эталона: keepLines + (опционально)
    keepNext, чтобы подпись не отрывалась от следующего элемента."""
    pPr = p._element.get_or_add_pPr()
    if keep_next and pPr.find(qn("w:keepNext")) is None:
        pPr.append(OxmlElement("w:keepNext"))
    if pPr.find(qn("w:keepLines")) is None:
        pPr.append(OxmlElement("w:keepLines"))
    pf = p.paragraph_format
    pf.alignment = alignment
    pf.first_line_indent = Cm(0)


def _make_caption_run(p, text):
    """Run подписи: bold, 12pt, color=auto (чёрный) — переопределение
    свойств стиля Caption (italic 9pt color text2) на уровне run, как
    сделано в эталоне отчёт_Курсовая.docx."""
    run = p.add_run(text)
    rPr = run._element.get_or_add_rPr()
    # bold
    b = OxmlElement("w:b")
    rPr.append(b)
    # color=auto — снимает темно-синий цвет стиля Caption
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "auto")
    rPr.append(color)
    # размер 12pt (sz=24 в полу-pt)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(szCs)
    return run


def add_figure_caption(doc, text):
    """Подпись под рисунком — по центру, bold, 12pt, чёрный.
    Стиль Caption (italic+keepLines), переопределяется bold/color/size
    на уровне run — как в эталоне."""
    p = doc.add_paragraph(style="Caption")
    _apply_caption_overrides(p, WD_ALIGN_PARAGRAPH.CENTER, keep_next=False)
    _make_caption_run(p, text)
    return p


def add_table_caption(doc, text):
    """Подпись над таблицей — по правому краю, bold, 12pt, чёрный,
    keepNext+keepLines, чтобы подпись не отрывалась от таблицы."""
    p = doc.add_paragraph(style="Caption")
    _apply_caption_overrides(p, WD_ALIGN_PARAGRAPH.RIGHT, keep_next=True)
    # spacing after = 0 (как в эталоне), чтобы подпись прижалась к таблице
    p.paragraph_format.space_after = Pt(0)
    _make_caption_run(p, text)
    return p


def new_document():
    """Создаёт новый Document на базе HSE-шаблона со всеми его стилями
    (HSE Default Text, HSE Title 1/2/3, HSE list, Caption, поля страницы,
    маркированный список)."""
    return Document(str(TEMPLATE_PATH))


def add_title_page(doc, part_subtitle, author_placeholder="________________________"):
    """Стандартный титул. На титуле явно задаём шрифт/размер run'ами, чтобы
    не зависеть от выбора стиля. После титула page break НЕ делаем — первый
    же HSE Title 1 сам начнётся с новой страницы (pageBreakBefore)."""
    # Удалим единственный пустой параграф, который был в шаблоне:
    body = doc.element.body
    first_p = body.find(qn("w:p"))
    if first_p is not None and not list(first_p.iter(qn("w:t"))):
        body.remove(first_p)

    def _centered(text, bold=False, italic=False, size=14):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic)
        return p

    _centered(
        "Пермский филиал федерального государственного автономного "
        "образовательного учреждения высшего образования\n"
        "«Национальный исследовательский университет\n"
        "«Высшая школа экономики»"
    )
    for _ in range(2):
        doc.add_paragraph()

    _centered("Факультет социально-экономических и компьютерных наук")

    for _ in range(3):
        doc.add_paragraph()

    _centered("РАЗРАБОТКА ВЕБ-СЕРВИСА ОПРЕДЕЛЕНИЯ ШКОЛЫ "
              "ПО АДРЕСУ ПРОЖИВАНИЯ НА ОСНОВЕ "
              "МУНИЦИПАЛЬНЫХ ПОСТАНОВЛЕНИЙ", bold=True)

    for _ in range(2):
        doc.add_paragraph()

    _centered("Отчёт о групповом проекте", italic=True)
    _centered(part_subtitle, italic=True, bold=True)

    doc.add_paragraph()

    _centered(
        "образовательной программы «Программная инженерия»\n"
        "по направлению подготовки 09.03.04 Программная инженерия"
    )

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Выполнил: {author_placeholder}")
    _set_font(run, size=14)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Руководитель: {author_placeholder}")
    _set_font(run, size=14)

    for _ in range(3):
        doc.add_paragraph()

    _centered("Пермь, 2026 год")
