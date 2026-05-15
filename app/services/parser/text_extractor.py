import shutil
import subprocess
import tempfile
from pathlib import Path

import pdfplumber
from docx import Document

from app.services.ocr.ocr_space import extract_text_with_ocr_space


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    lines = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_text_from_pdf_native(file_path: str) -> str:
    lines = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                lines.append(text)

    return "\n\n".join(lines)


def extract_text_from_pdf(file_path: str) -> str:
    native_text = extract_text_from_pdf_native(file_path)

    if len(native_text.strip()) >= 500:
        return native_text

    return extract_text_with_ocr_space(file_path)


def convert_doc_to_docx(file_path: str) -> str | None:
    """
    Конвертация .doc → .docx через LibreOffice (`soffice`). Кросс-платформенно.
    Возвращает путь к получившемуся .docx, либо None если LibreOffice не
    установлен или конвертация упала.
    """
    if not shutil.which("soffice"):
        return None

    tmp_dir = tempfile.mkdtemp(prefix="doc_convert_")

    try:
        subprocess.run(
            [
                "soffice", "--headless",
                "--convert-to", "docx",
                "--outdir", tmp_dir,
                file_path,
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=180,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return None

    converted = list(Path(tmp_dir).glob("*.docx"))
    if not converted:
        return None

    return str(converted[0])


def extract_text_from_doc(file_path: str) -> str:
    """
    Сначала пробуем кросс-платформенный путь через LibreOffice
    (.doc → .docx → текст). Если LibreOffice не установлен,
    падаем на macOS-only `textutil`.
    """
    converted = convert_doc_to_docx(file_path)
    if converted:
        return extract_text_from_docx(converted)

    if shutil.which("textutil"):
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", file_path],
            capture_output=True,
            text=True,
            check=True,
        )
        return result.stdout

    raise RuntimeError(
        "Не удалось извлечь текст из .doc: установите LibreOffice "
        "(`brew install libreoffice` / `apt install libreoffice-writer`) "
        "или конвертируйте файл в .docx / .pdf вручную."
    )


def extract_text(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()

    if suffix == ".docx":
        return extract_text_from_docx(file_path)

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)

    if suffix == ".doc":
        return extract_text_from_doc(file_path)

    raise ValueError(f"Неподдерживаемый формат файла: {suffix}")
