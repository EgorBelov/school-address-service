"""
Извлечение правил из «многоколоночных» таблиц вида:

    №  |  Наименование и адрес школы  |  Улица  |  Номер дома
    1  |  МАОУ «Гимназия № 31» ...    |  Блюхера         |  3, 5, 7, 7а
    1  |  МАОУ «Гимназия № 31» ...    |  Вильвенская     |  1, 2, 3, 5, 6
    2  |  МАОУ «СОШ № 50» ...         |  Подлесная       |  43а

Каждая строка таблицы — это ОДНО правило для одной школы. Имя школы
часто повторяется в каждой строке (или указывается только в первой
строке группы, а остальные имеют пустую ячейку).

Формат типичен для постановлений Пермского края, Перми, и т.д.
"""
import re
from pathlib import Path

import pdfplumber
from docx import Document

from app.services.parser.pdf_table_extractor import (
    SCHOOL_NAME_RE,
    SCHOOL_PREFIX_RE,
    clean,
)


_HEADER_MARKERS = (
    "наименование",
    "номер дома",
    "номер\nдома",
    "территория",
    "адрес",
    "улица",
    "п/п",
)


def _looks_like_header(row: list) -> bool:
    """Шапка таблицы: содержит слова «Наименование», «улица», «номер дома» и т.п."""
    text = " ".join(str(c or "").lower() for c in row)
    if not text.strip():
        return False
    marker_hits = sum(1 for m in _HEADER_MARKERS if m in text)
    return marker_hits >= 2


def _is_school_or_prefix(text: str) -> bool:
    if not text:
        return False
    stripped = text.strip()
    if SCHOOL_NAME_RE.search(stripped):
        return True
    if SCHOOL_PREFIX_RE.match(stripped):
        return True
    return False


def _has_house_numbers(text: str) -> bool:
    """В ячейке «Номер дома» должна быть хоть одна цифра."""
    return bool(re.search(r"\d", text or ""))


def _is_street_like(text: str) -> bool:
    """В ячейке «Улица» — название улицы; обычно без цифр, длина 3–60."""
    if not text:
        return False
    text = text.strip()
    if len(text) < 2 or len(text) > 80:
        return False
    if text.lower() in ("улица", "ул.", "район", "—", "-"):
        return False
    return True


def split_street_aliases(street: str) -> list[str]:
    """
    Названия в скобках (старое/альтернативное название улицы) — это
    тот же геообъект, и поиск должен находить и по новому, и по
    старому имени:
        «Монастырская (Орджоникидзе)» → ['Монастырская', 'Орджоникидзе']
        «Парковый проспект»            → ['Парковый проспект']
    """
    if not street:
        return []

    match = re.match(r"^(.+?)\s*\(([^)]+)\)\s*$", street)
    if not match:
        return [street.strip()]

    main = match.group(1).strip()
    alias = match.group(2).strip()

    result = []
    if main:
        result.append(main)
    if alias and alias.lower() != main.lower():
        result.append(alias)
    return result


def extract_multi_column_rows_from_table(
    table: list[list],
    page_index: int,
    current_school: str | None = None,
) -> tuple[list[dict], str | None]:
    """
    Преобразует одну таблицу (4+ колонок) в raw-правила.
    Возвращает (правила, current_school) — школу нужно пробросить
    в следующий вызов: большая таблица постановления может быть
    разбита pdfplumber'ом на отдельные таблицы по страницам, и
    continuation-строки на новой странице должны привязаться к
    школе, начатой на предыдущей.
    """
    rows: list[dict] = []

    for raw in table:
        if not raw or len(raw) < 4:
            continue

        # Берём только первые 4 колонки (бывает шире — игнорируем)
        cells = [clean(c) for c in raw[:4]]
        _, school_col, street_col, houses_col = cells

        if _looks_like_header(cells):
            continue

        # Обновляем имя школы, если в этой строке оно валидно
        if _is_school_or_prefix(school_col):
            current_school = school_col

        if not current_school:
            continue

        if not _is_street_like(street_col):
            continue

        if not _has_house_numbers(houses_col):
            continue

        for street_name in split_street_aliases(street_col):
            rows.append({
                "page": page_index,
                "school_name": current_school,
                "street": street_name,
                "houses": houses_col,
            })

    return rows, current_school


def extract_multi_column_rows_from_pdf(file_path: str) -> list[dict]:
    rows: list[dict] = []
    current_school: str | None = None

    with pdfplumber.open(file_path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            for table in page.extract_tables() or []:
                if not table or not table[0] or len(table[0]) < 4:
                    continue
                new_rows, current_school = extract_multi_column_rows_from_table(
                    table, page_index, current_school
                )
                rows.extend(new_rows)

    print(f"[multi_column_extractor] PDF rows: {len(rows)}")
    return rows


def extract_multi_column_rows_from_docx(file_path: str) -> list[dict]:
    doc = Document(file_path)
    rows: list[dict] = []
    current_school: str | None = None

    for table_index, table in enumerate(doc.tables, start=1):
        if not table.rows or len(table.rows[0].cells) < 4:
            continue
        raw_table = [[c.text for c in r.cells] for r in table.rows]
        new_rows, current_school = extract_multi_column_rows_from_table(
            raw_table, table_index, current_school
        )
        rows.extend(new_rows)

    print(f"[multi_column_extractor] DOCX rows: {len(rows)}")
    return rows


def _row_looks_like_multi_data(cells: list[str]) -> bool:
    """
    Строка-данные многоколоночного формата: в col[0] — номер пункта,
    в col[1] — название школы. Используется когда в файле нет
    нормального заголовка таблицы (часто бывает в DOCX, где каждая
    школа = отдельная таблица без header).
    """
    if len(cells) < 4:
        return False
    n_col = (cells[0] or "").strip()
    school_col = (cells[1] or "").strip()
    if not n_col or not school_col:
        return False
    if not re.match(r"^\d+(\.\d+)?$", n_col):
        return False
    return _is_school_or_prefix(school_col)


def is_multi_column_format(file_path: str) -> bool:
    """
    Структурная эвристика, без LLM. Распознаём два сценария:
      A. Таблица ≥4 колонок с шапкой «Наименование … | улица | номер дома».
      B. ≥3 таблиц по 4 колонки, у которых в col[0] стоит номер пункта,
         а в col[1] — название школы (DOCX Пермского края: каждая школа —
         собственная мини-таблица).
    """
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        try:
            with pdfplumber.open(file_path) as pdf:
                data_table_hits = 0
                for page in pdf.pages[:15]:
                    for table in page.extract_tables() or []:
                        if not table or not table[0]:
                            continue
                        if len(table[0]) < 4:
                            continue
                        # Сценарий A
                        if any(_looks_like_header(r) for r in table[:3]):
                            return True
                        # Сценарий B
                        for row in table[:5]:
                            if _row_looks_like_multi_data([clean(c) for c in row[:4]]):
                                data_table_hits += 1
                                break
                if data_table_hits >= 3:
                    return True
        except Exception:
            return False

    if suffix == ".docx":
        try:
            doc = Document(file_path)
            data_table_hits = 0
            for table in doc.tables[:50]:
                if not table.rows:
                    continue
                cells = [c.text for c in table.rows[0].cells]
                if len(cells) < 4:
                    continue
                # Сценарий A
                if _looks_like_header(cells):
                    return True
                # Сценарий B
                for row in table.rows[:5]:
                    row_cells = [clean(c.text) for c in row.cells[:4]]
                    if _row_looks_like_multi_data(row_cells):
                        data_table_hits += 1
                        break
            if data_table_hits >= 3:
                return True
        except Exception:
            return False

    return False


# ────────────────── превращение в структуру decree → schools[] → rules[] ─────

def _split_houses(houses: str) -> list[str]:
    """«1, 2, 3, 5а, 7-9» → ['1', '2', '3', '5а', '7-9']."""
    if not houses:
        return []
    return [item.strip() for item in re.split(r"[;,]", houses) if item.strip()]


def _looks_like_range(item: str) -> bool:
    return bool(re.match(r"^\s*\d+\s*[-–]\s*\d+\s*$", item))


def _parse_range(item: str) -> tuple[int, int] | None:
    m = re.match(r"^\s*(\d+)\s*[-–]\s*(\d+)\s*$", item)
    if not m:
        return None
    return int(m.group(1)), int(m.group(2))


def rows_to_decree_dict(rows: list[dict], metadata: dict) -> dict:
    """
    Группирует raw-строки по школам и собирает финальную структуру
    {decree, schools: [{name, address, rules: []}]}.
    """
    by_school: dict[str, dict] = {}

    for row in rows:
        name = row["school_name"]

        if name not in by_school:
            by_school[name] = {
                "name": name,
                "address": "",
                "rules": [],
            }

        houses_raw = row["houses"]
        items = _split_houses(houses_raw)

        # Если все элементы — отдельные дома (числа/числа-с-буквами),
        # формируем exact_list. Если есть диапазоны — мердж их в
        # отдельные rules с house_from/house_to.
        exact_items: list[str] = []
        range_items: list[tuple[int, int]] = []
        for item in items:
            if _looks_like_range(item):
                rng = _parse_range(item)
                if rng:
                    range_items.append(rng)
                else:
                    exact_items.append(item)
            else:
                exact_items.append(item)

        if exact_items:
            by_school[name]["rules"].append({
                "locality": "",
                "street": row["street"],
                "house_rule_raw": houses_raw,
                "parity": "all",
                "house_from": None,
                "house_to": None,
                "house_number": ",".join(exact_items),
                "house_numbers": exact_items,
                "rule_type": "exact_list",
                "comment": "",
            })

        for h_from, h_to in range_items:
            by_school[name]["rules"].append({
                "locality": "",
                "street": row["street"],
                "house_rule_raw": f"{h_from}-{h_to}",
                "parity": "all",
                "house_from": h_from,
                "house_to": h_to,
                "house_number": None,
                "house_numbers": [],
                "rule_type": "range",
                "comment": "",
            })

    return {
        "decree": {
            "number": metadata.get("number", ""),
            "date": metadata.get("date", ""),
            "municipality": metadata.get("municipality", ""),
        },
        "schools": list(by_school.values()),
        "errors": [],
    }
